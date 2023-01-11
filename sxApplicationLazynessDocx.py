from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from os.path import exists
import os
from time import sleep
import re
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH

if(exists("applications.docx")):
	os.remove("applications.docx")

pages = int(input("Hvor mange siders ansøgninger er der? "))
teams = int(input("Hvor mange hold er der? "))

driver = webdriver.Chrome(service = Service(ChromeDriverManager().install()))
applications = []
def printProgressBar (iteration, total, prefix = '', suffix = '', decimals = 1, length = 100, fill = '█', printEnd = "\r"):
    """
    Call in a loop to create terminal progress bar
    @params:
        iteration   - Required  : current iteration (Int)
        total       - Required  : total iterations (Int)
        prefix      - Optional  : prefix string (Str)
        suffix      - Optional  : suffix string (Str)
        decimals    - Optional  : positive number of decimals in percent complete (Int)
        length      - Optional  : character length of bar (Int)
        fill        - Optional  : bar fill character (Str)
        printEnd    - Optional  : end character (e.g. "\r", "\r\n") (Str)
    """
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filledLength = int(length * iteration // total)
    bar = fill * filledLength + '-' * (length - filledLength)
    print(f'\r{prefix} |{bar}| {percent}% {suffix}', end = printEnd)
    # Print New Line on Complete
    if iteration == total: 
        print()

def add_hyperlink(paragraph, text, url):
	# This gets access to the document.xml.rels file and gets a new relation id value
	part = paragraph.part
	r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

	# Create the w:hyperlink tag and add needed values
	hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
	hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

	# Create a w:r element and a new w:rPr element
	new_run = docx.oxml.shared.OxmlElement('w:r')
	rPr = docx.oxml.shared.OxmlElement('w:rPr')

	# Join all the xml elements together add add the required text to the w:r element
	new_run.append(rPr)
	new_run.text = text
	hyperlink.append(new_run)

	# Create a new Run object and add the hyperlink into it
	r = paragraph.add_run ()
	r._r.append (hyperlink)

	# A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
	# Delete this if using a template that has the hyperlink style in it
	r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
	r.font.underline = True
	r.font.name = "Arial"

	return hyperlink

def addHeading(text, document):
	heading = document.add_heading("", 2)
	run = heading.add_run(text)
	run.bold = True
	run.font.name = "Arial"

def addApplicantInformation(text, tableRow):
	p = tableRow.add_paragraph()
	run = p.add_run(text)
	run.bold = True
	run.font.name = "Arial"
	p.add_run().add_break()

def addApplicantInformationWithLink(text, linkText, Link, tableRow):
	p = tableRow.add_paragraph()
	run = p.add_run(text)
	run.bold = True
	run.font.name = "Arial"
	add_hyperlink(p, linkText, Link)

def formatter(applications, teams):
	document = Document()
	print("Creating document")
	applications = sorted(applications, key=lambda d: d['name']) 
	for i in range(1, int(teams) + 1):
		addHeading("HOLD "+str(i)+": ", document)
	
	overViewTable = document.add_table(rows = 1, cols = 3)
	overViewTable.style = 'Table Grid'
	overViewTableHeads = overViewTable.rows[0].cells
	
	run = overViewTableHeads[0].paragraphs[0].add_run("Navn")
	run.bold = True
	run.font.name = "Arial"

	run = overViewTableHeads[1].paragraphs[0].add_run("Er blevet snakket med? ✅❌")
	run.bold = True
	run.font.name = "Arial"

	run = overViewTableHeads[2].paragraphs[0].add_run("Godkendt/afvist ✅❌🤷")
	run.bold = True
	run.font.name = "Arial"

	for application in applications:
		row_cells = overViewTable.add_row().cells
		run = row_cells[0].paragraphs[0].add_run(application["name"])
		run.font.name = "Arial"
		run = row_cells[1].paragraphs[0].add_run("❌")
		run.font.name = "Arial"
		row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	document.add_page_break()

	for i, application in enumerate(applications):
		applicantTable = document.add_table(rows = 1, cols = 1)
		applicantTable.style = 'Table Grid'
		applicantTableHeading = applicantTable.rows[0].cells
		run = applicantTableHeading[0].paragraphs[0].add_run(application["name"])
		run.bold = True
		run.font.name = "Arial"
		applicantTableHeading[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

		applicationRow = applicantTable.add_row().cells[0]
		
		addApplicantInformationWithLink("Navn: ", application["name"], application["link"], applicationRow)
		addApplicantInformationWithLink("SteamID: ", application["steamid"], application["dashboard"], applicationRow)
		applicationRow.add_paragraph()
		addApplicantInformation("Godt: ", applicationRow)
		addApplicantInformation("Dårligt: ", applicationRow)
		addApplicantInformation("Andet: ", applicationRow)
		addApplicantInformation("Beslutning: ", applicationRow)
		addApplicantInformation("Grund (Hvis afvist): ", applicationRow)
		
		document.add_paragraph()
		printProgressBar(i + 1, l, prefix = 'Progress:', suffix = 'Complete', length = 50)
	print("Saving document")
	document.save('applications.docx')

def fetchApplications():
	print("Fetching applications")
	for i in range(1, pages + 1):
		driver.get("https://stavox.dk/forums/forum/5-staffans%C3%B8gninger-%C3%A5bne/?page=" + str(i))
		list1 = driver.find_element(By.XPATH,'/html/body/main/div/div/div/div[3]/div/ol')
		elements = list1.find_elements(By.CSS_SELECTOR, "li > div > h4 > span > a")
		sleep(2)
		for element in elements:
			if re.search("\[Accepteret\]", element.text) == None and re.search("\[Afslået\]", element.text) == None:
				if (element.text != "Mundtlig Staffansøgning - How to [BEMÆRK NY FORMAT]"):
					applicant = {
						"name": element.text.title(),
						"link": element.get_attribute('href'),
						"steamid": "",
						"dashboard": "",
					}
					applications.append(applicant)

fetchApplications()
numOfApplications = 0
l = len(applications)
printProgressBar(0, l, prefix = 'Progress:', suffix = 'Complete', length = 50)
for i, application in enumerate(applications):
	driver.get(application["link"])
	sleep(0.5)
	steamid = WebDriverWait(driver, 10).until(
		EC.presence_of_element_located((By.XPATH, '/html/body/main/div/div/div/div[3]/div[2]/form/article[1]'))
	)
	lowerApplicationText = steamid.text.lower()
	if(("steamid" in lowerApplicationText) or ("steam id" in lowerApplicationText)):
		if(re.search("STEAM_.*", steamid.text)):
			applicantSteamID = re.findall("STEAM_.*", steamid.text)
		else:
			applicantSteamID = re.findall("steam_.*", steamid.text)
		application['steamid'] = applicantSteamID[0]
		application['dashboard'] = "https://stavox.dk/dash/lookup?lookupid=" + applicantSteamID[0]
	numOfApplications = numOfApplications + 1
	printProgressBar(i + 1, l, prefix = 'Progress:', suffix = 'Complete', length = 50)

printProgressBar(0, l, prefix = 'Progress:', suffix = 'Complete', length = 50)
formatter(applications, teams)
print(f'Found a total of {numOfApplications} applications')
print("Applications has been saved in applications.docx")